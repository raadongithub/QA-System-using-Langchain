import os
import streamlit as st
from streamlit_chat import message
from dotenv import load_dotenv
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import google.generativeai as genai
from google.generativeai import GenerativeModel
import io
from PyPDF2 import PdfReader  # PDF handling
from docx import Document as DocxDocument  # DOCX handling


# Class for handling document retrieval and chat interaction
class SimpleConversationalRetriever:
    def __init__(self, retriever, model):
        self.retriever = retriever
        self.model = model
        self.chat_history = []

    def ask(self, question):
        docs = self.retriever.get_relevant_documents(question)
        context = "\n".join([doc.page_content for doc in docs])
        prompt = f"{context}\n\nQuestion: {question}\nAnswer:"
        response = self.model.generate_content(prompt)
        answer = response.text.strip()

        self.chat_history.append({"question": question, "answer": answer})
        return answer


# Configure Google Gemini API with user's API key
def configure_gemini():
    api_key = st.sidebar.text_input("Google API Key:", type="password")
    if not api_key:
        st.sidebar.warning("Please enter your Google API Key.")
        return None
    genai.configure(api_key=api_key)
    return GenerativeModel("gemini-pro")


# Load documents (PDF, DOCX, or TXT) based on file type
def load_documents(file):
    name, ext = os.path.splitext(file.name)

    if ext == ".pdf":
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return [Document(page_content=text)]

    elif ext == ".docx":
        docx = DocxDocument(file)
        text = ""
        for para in docx.paragraphs:
            text += para.text + "\n"
        return [Document(page_content=text)]

    elif ext == ".txt":
        text_data = file.read().decode("utf-8")
        return [Document(page_content=text_data)]

    else:
        st.write(f"Document format {ext} is not supported!")
        return None


# Split document text into chunks for processing
def chunk_data(data):
    if not data:
        st.error("No data found in the file.")
        return []

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    chunks = text_splitter.split_documents(data)
    return chunks


# Create a vector store for embedding chunks using HuggingFace embeddings
def create_vector_store(chunks):
    if not chunks:
        st.error("No chunks to embed.")
        return None

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
    )
    vector_store = FAISS.from_documents(chunks, embeddings)
    return vector_store


def main():
    st.set_page_config(
        page_title="ASR Question-Answering Application", page_icon=":books:"
    )
    st.subheader("ASR Question-Answering Application 🤖")

    with st.expander("About the App"):
        st.markdown(
            """
            This Chatbot will respond to the queries that you have for the document uploaded.
            """
        )

    # Configure the model
    gemini_model = configure_gemini()
    if not gemini_model:
        return

    # Upload and process the file
    uploaded_file = st.sidebar.file_uploader(
        "Upload a file:", type=["pdf", "docx", "txt"]
    )
    if uploaded_file:

        with st.spinner("Processing"):
            file_data = load_documents(uploaded_file)
            if file_data:
                chunks = chunk_data(file_data)
                vector_store = create_vector_store(chunks)
                if vector_store:
                    st.success("File uploaded, chunked, and embedded successfully")

                    retriever = vector_store.as_retriever()
                    conversation_handler = SimpleConversationalRetriever(
                        retriever, gemini_model
                    )

                    # Initialize chat history
                    if "history" not in st.session_state:
                        st.session_state["history"] = []

                    if "generated" not in st.session_state:
                        st.session_state["generated"] = [
                            "Hello! Ask me anything about " + uploaded_file.name + " 🤗"
                        ]

                    if "past" not in st.session_state:
                        st.session_state["past"] = ["Hey! 👋"]

                    # Containers for chat and input
                    response_container = st.container()
                    container = st.container()

                    # User input form
                    with container:
                        with st.form(key="my_form", clear_on_submit=True):
                            user_input = st.text_input(
                                "Query:",
                                placeholder="Ask about your document (:",
                                key="input",
                            )
                            submit_button = st.form_submit_button(label="Send")

                        if submit_button and user_input:
                            output = conversation_handler.ask(user_input)

                            st.session_state["past"].append(user_input)
                            st.session_state["generated"].append(output)

                    # Display the conversation history
                    if st.session_state["generated"]:
                        with response_container:
                            for i in range(len(st.session_state["generated"])):
                                message(
                                    st.session_state["past"][i],
                                    is_user=True,
                                    key=str(i) + "_user",
                                )
                                message(st.session_state["generated"][i], key=str(i))
                else:
                    st.error("Failed to create a vector store from the file.")
            else:
                st.error("Failed to load data from the file.")


if __name__ == "__main__":
    main()
